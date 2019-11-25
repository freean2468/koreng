import scrapy
import re
from docx import Document

class ScriptSpider(scrapy.Spider):
    name = "scripts"
    #'https://www.springfieldspringfield.co.uk/'
    start_urls = [
        'https://www.springfieldspringfield.co.uk/view_episode_scripts.php?tv-show=spartacus-gods-of-the-arena&episode=s01e01'
    ]

    def __init__(self):
        self.path = '/Users/hoon-ilsong/Documents/crawler products/'

    def parse(self, response):
        page = response.url.split("/")[-2]
        document = Document()

        # u'The Vampire Diaries s04e15 Episode Script            '
        h = response.css("div.main-content-left h1::text").get()
        title = re.sub("[']", '', h).strip()
        # title = title.replace('u', '', 1)

        subtitle = response.css("div.main-content-left h3::text").get()
        document.add_heading(subtitle, 0)

        res = [i for i in range(len(title)) if title.startswith(' ', i)]
        title = title[0:res[-2]];

        body = response.css("div.scrolling-script-container::text").getall()

        document.add_paragraph(body)
        document.save(self.path+title+'.docx')

        self.log('<------------------------------ Saved file %s ----------------------------->' % title)

        next_page = ''

        i = 0
        PorN = response.css("div.episode_script a::text").getall()
        while(i < len(PorN)):
            _temp = PorN[i]
            _temp = re.sub("[']", '', _temp)
            self.log(_temp)
            if _temp == 'Next Episode':
                next_page = response.css("div.episode_script a::attr(href)").getall()[i]
                break
            i=i+1

        if next_page is not None:
            next_page = response.urljoin(next_page)
            yield scrapy.Request(next_page, callback=self.parse)
